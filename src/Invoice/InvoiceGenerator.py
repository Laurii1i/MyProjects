import os
from docx import Document
from docx.shared import Pt
from copy import deepcopy
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from UI.BlackHole import BlackHole
PATH = os.path.realpath(__file__)
PATH = os.path.dirname(PATH)

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "color": "#0000FF"},
        end={"sz": 12, "val": "dashed", "color": "#000000"}
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

class InvoiceGeneratorWord:

    def __init__(self, blackholelist)  -> None:
         
        self.doc = Document(os.path.join(PATH,'laskupohja.docx'))
        self.black_holes = blackholelist
        self.table = self.doc.tables[1]
        self.template_row = self.table.rows[1]
        self.table_height = self.calculate_table_height()
        self.row_heigth = self.table.rows[1].height
        self.table.rows[2].height = self.row_heigth 
        self.veroton_hinta_yht = 0.0
        for cell in self.table.rows[2].cells:
            set_cell_border(
                cell,
                top={"sz": 6, "val": "single", "color": "#000000"},
                bottom={"sz": 0, "val": "none"},
                start={"sz": 6, "val": "single", "color": "#000000"},
                end={"sz": 6, "val": "single", "color": "#000000"}
            )


    def generate_invoice(self):
        cell = self.table.cell(1,0)
        self.set_text(cell, 'Työt_muo')
        for black_hole in self.black_holes:
            self.add_kori(black_hole)
        self.remove_template_lines()
        self.resize_table_to_original_size()
        self.update_summary()
        self.doc.save(os.path.join(PATH,'laskupohja_muokattu.docx'))

    def add_kori(self, black_hole: BlackHole):
        self.add_kori_line(black_hole.basket_frame.title_label._text)
        for image in black_hole.basket_frame.images:
            self.add_image_line(image)

    def add_kori_line(self, string : str):
        new_row = self.add_row()
        cell = new_row.cells[0]
        self.copy_cell_format(self.template_row.cells[0], cell) 
        self.set_text(cell, string) 

    def add_kori_line(self, string : str):
        new_row = self.add_row()
        cell = new_row.cells[0]
        self.set_text(cell, string)

    def add_image_line(self, image):
        new_row = self.add_row()
        # Kuvaus
        cell = new_row.cells[0] 
        
        self.set_text(cell, image.product_data['name'] + ', ' + image.product_data['color'])
        # Määrä
        cell = new_row.cells[1]
        self.set_text(cell, '1')
        # Yksikkö
        cell = new_row.cells[2]
        self.set_text(cell, 'kpl')
        # A-hinta
        cell = new_row.cells[3]
        self.set_text(cell, image.product_data['price'] + ' €')
        # Alennus
        cell = new_row.cells[4]
        self.set_text(cell, str(image.discount)+' %')
        # Yhteensä
        cell = new_row.cells[5]
        hinta = float(image.product_data['price'])*(1-float(image.discount)/100)
        self.veroton_hinta_yht = self.veroton_hinta_yht + hinta 
        self.set_text(cell, str(hinta))

    def add_row(self):
        new_row = self.table.add_row()
        new_row.height = self.row_heigth
        for i in range(len(new_row.cells)):
            cell = new_row.cells[i]
            self.copy_cell_format(self.template_row.cells[i], cell) 
            self.set_text(cell, '')
            set_cell_border(
                cell,
                top={"sz": 6, "val": "single", "color": "#000000"},
                bottom={"sz": 6, "val": "single", "color": "#000000"},
                start={"sz": 6, "val": "single", "color": "#000000"},
                end={"sz": 6, "val": "single", "color": "#000000"}
            )
        return new_row
    
    def calculate_table_height(self):
        table_height = 0
        self.table.rows
        for i in range(len(self.table.rows)):
            table_height = table_height + self.table.rows[i].height
        return table_height

    def set_text(self, cell, string: str):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ''
        # add_paragraph
        if len(cell.paragraphs) > 0:
            if len(cell.paragraphs[0].runs) > 0:
                cell.paragraphs[0].runs[0].text = ' '+string
            else:
                cell.paragraphs[0].add_run('asdasd')
    
    def add_line(self):
        new_row = self.table.add_row()
        new_row.height = self.row_heigth
        for i in range(len(new_row.cells)):
            cell = new_row.cells[i]
            self.copy_cell_format(self.template_row.cells[i], cell) 
            self.set_text(cell, 'new_line')
            set_cell_border(
                cell,
                top={"sz": 6, "val": "single", "color": "#000000"},
                bottom={"sz": 6, "val": "single", "color": "#000000"},
                start={"sz": 6, "val": "single", "color": "#000000"},
                end={"sz": 6, "val": "single", "color": "#000000"}
            )
    
    def copy_cell_format(self, source_cell, target_cell):
        # Copy paragraph formats
        target_cell.paragraphs[0].paragraph_format._element = deepcopy(source_cell.paragraphs[0].paragraph_format._element)
    
        # Clear existing content while preserving formatting
        for paragraph in target_cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        
        # Ensure there's at least one run in the target cell
        if not target_cell.paragraphs[0].runs:
            target_cell.paragraphs[0].add_run()
        
        # Copy run properties
        if source_cell.paragraphs[0].runs:
            source_run = source_cell.paragraphs[0].runs[0]
            target_run = target_cell.paragraphs[0].runs[0]
            
            # Copy font properties
            target_run.font.name = source_run.font.name
            target_run.font.size = source_run.font.size
            target_run.font.bold = source_run.font.bold
            target_run.font.italic = source_run.font.italic
            
            # Copy other run properties
            if source_run._element.rPr is not None:
                target_run._element.get_or_add_rPr().append(deepcopy(source_run._element.rPr))

    def remove_template_lines(self):
        def remove_table_row(table, row_index):
            tbl = table._tbl
            tr = tbl.tr_lst[row_index]
            tbl.remove(tr)

        remove_table_row(self.table, 1)       
        remove_table_row(self.table, 1) 

    def resize_table_to_original_size(self):
        current_table_height = self.calculate_table_height()
        height_to_be_added = self.table_height - current_table_height
        self.table.rows[-1].height = self.table.rows[-1].height + height_to_be_added
    
    def update_summary(self):
        veroton_hinta = self.veroton_hinta_yht
        self.doc.paragraphs[6].text = f'Veroton hinta\t{veroton_hinta} €'
        self.doc.paragraphs[7].text = f'Arvonlisävero\t{veroton_hinta*0.24:.2f} €'
        self.doc.paragraphs[10].text = f'Yhteensä\t{veroton_hinta*1.24:.2f} €'


if __name__ == '__main__':
    PATH = os.path.realpath(__file__)
    PATH = os.path.dirname(PATH)
    g = InvoiceGeneratorWord()
    g.generate_invoice()